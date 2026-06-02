"""
Convert a Markdown file to a formatted .docx.

Handles the subset of Markdown actually used in our docs:
  - # / ## / ### headings
  - **bold** and *italic* inline
  - bulleted lists (- or *) and numbered lists (1.)
  - tables ( | col | col | with ---- separators )
  - fenced code blocks (``` … ```)
  - inline `code`
  - blockquotes (> )
  - horizontal rules (---)
  - links: rendered as "label (URL)"

Usage:
    python scripts/md_to_docx.py <input.md> <output.docx>
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


INLINE_PATTERN = re.compile(
    r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\))"
)
LINK_PATTERN = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def add_inline_runs(paragraph, text: str) -> None:
    """Walk a line of markdown, splitting into runs with the right styles."""
    pos = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            r = paragraph.add_run(token[2:-2])
            r.bold = True
        elif token.startswith("*") and token.endswith("*"):
            r = paragraph.add_run(token[1:-1])
            r.italic = True
        elif token.startswith("`") and token.endswith("`"):
            r = paragraph.add_run(token[1:-1])
            r.font.name = "Consolas"
            r.font.color.rgb = RGBColor(0x66, 0x33, 0x99)
        elif token.startswith("["):
            m = LINK_PATTERN.match(token)
            if m:
                label, url = m.group(1), m.group(2)
                r = paragraph.add_run(f"{label} ")
                r2 = paragraph.add_run(f"({url})")
                r2.font.size = Pt(9)
                r2.font.color.rgb = RGBColor(0x33, 0x66, 0xCC)
            else:
                paragraph.add_run(token)
        else:
            paragraph.add_run(token)
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and s.count("|") >= 2


def parse_table_cells(line: str) -> list[str]:
    s = line.strip()
    # strip leading/trailing pipes, split on |
    inner = s[1:-1] if s.startswith("|") and s.endswith("|") else s
    return [c.strip() for c in inner.split("|")]


def convert(src_path: Path, out_path: Path) -> None:
    text = src_path.read_text(encoding="utf-8")
    lines = text.split("\n")

    doc = Document()
    # Reasonable default paragraph font for body
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    i = 0
    in_code = False
    code_buf: list[str] = []

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        # Fenced code blocks
        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                # Close: flush as a single monospace paragraph block
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buf))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                p.paragraph_format.left_indent = Pt(12)
                in_code = False
                code_buf = []
            i += 1
            continue

        if in_code:
            code_buf.append(raw)
            i += 1
            continue

        # Tables — consume a contiguous block of |…| lines
        if is_table_row(line):
            header_cells = parse_table_cells(line)
            i += 1
            # The next line should be a --- separator; consume if so
            if i < len(lines) and re.match(r"^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$", lines[i]):
                i += 1
            rows: list[list[str]] = []
            while i < len(lines) and is_table_row(lines[i]):
                rows.append(parse_table_cells(lines[i]))
                i += 1

            t = doc.add_table(rows=1 + len(rows), cols=len(header_cells))
            t.style = "Light Grid Accent 1"
            hdr = t.rows[0].cells
            for col_idx, cell_text in enumerate(header_cells):
                cell = hdr[col_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                add_inline_runs(p, cell_text)
                for r in p.runs:
                    r.bold = True
            for row_idx, row in enumerate(rows, start=1):
                for col_idx, cell_text in enumerate(row):
                    if col_idx >= len(t.rows[row_idx].cells):
                        continue
                    cell = t.rows[row_idx].cells[col_idx]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    add_inline_runs(p, cell_text)
            doc.add_paragraph()
            continue

        # Horizontal rule
        if line.strip() == "---":
            p = doc.add_paragraph()
            run = p.add_run("─" * 60)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # Headings
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue
        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
            i += 1
            continue

        # Blockquote
        if line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            add_inline_runs(p, line[2:].strip())
            for r in p.runs:
                r.italic = True
            i += 1
            continue

        # Bulleted list
        bullet_match = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if bullet_match:
            indent_spaces = len(bullet_match.group(1))
            content = bullet_match.group(2)
            level = min(indent_spaces // 2, 3)
            style = "List Bullet" if level == 0 else f"List Bullet {level + 1}"
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, content)
            i += 1
            continue

        # Numbered list
        number_match = re.match(r"^(\s*)\d+\.\s+(.*)$", line)
        if number_match:
            content = number_match.group(2)
            try:
                p = doc.add_paragraph(style="List Number")
            except KeyError:
                p = doc.add_paragraph()
            add_inline_runs(p, content)
            i += 1
            continue

        # Blank line
        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue

        # Plain paragraph
        p = doc.add_paragraph()
        add_inline_runs(p, line)
        i += 1

    doc.save(out_path)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("usage: python scripts/md_to_docx.py <input.md> <output.docx>")
        sys.exit(1)
    convert(Path(sys.argv[1]), Path(sys.argv[2]))
    print(f"Wrote {sys.argv[2]}")
